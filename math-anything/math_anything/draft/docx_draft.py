"""Convert Markdown methodology draft to Word (.docx) document."""

from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def markdown_to_docx(md_text: str, output_path: str, title: str = "Computational Methodology") -> None:
    """Convert a Markdown methodology draft to a .docx file.

    Args:
        md_text: Markdown-formatted methodology text.
        output_path: Path to write the .docx file.
        title: Document title.
    """
    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Simple Markdown parser
    lines = md_text.splitlines()
    in_code = False
    buffer: List[str] = []
    code_buffer: List[str] = []

    def flush_buffer() -> None:
        if not buffer:
            return
        text = " ".join(buffer).strip()
        if text:
            doc.add_paragraph(text)
        buffer.clear()

    for line in lines:
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code:
                flush_buffer()
                # Add buffered code as a formatted paragraph
                code_text = "\n".join(code_buffer)
                if code_text:
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_buffer.clear()
                in_code = False
            else:
                flush_buffer()
                code_buffer.clear()
                in_code = True
            continue

        if in_code:
            code_buffer.append(stripped)
            continue

        # Headings
        if stripped.startswith("# "):
            flush_buffer()
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            flush_buffer()
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            flush_buffer()
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            flush_buffer()
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("1. ") or stripped.startswith("2. ") or stripped.startswith("3. "):
            flush_buffer()
            p = doc.add_paragraph(stripped[3:], style="List Number")
        elif stripped == "":
            flush_buffer()
        else:
            buffer.append(line)

    flush_buffer()
    doc.save(output_path)
