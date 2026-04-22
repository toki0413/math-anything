"""
CBES论文结构优化器
基于数学分析结果重新优化论文结构

核心优化目标：
1. 打破SCC-2循环依赖 (Mech↔Batt↔Sup→单向流)
2. 压缩至450段 (精简27%冗余)
3. 添加Mech→Sup桥梁段落
4. 优化可读性 (FK指数从22.4/21.6降至<20)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os

# 数学分析驱动的优化策略
OPTIMIZATION_STRATEGY = {
    # 1. 打破循环依赖：重新定义章节流向
    "dependency_flow": {
        "old": "Intro→Mech↔Batt↔Sup↔Fab→App→Eco→Con",  # 循环
        "new": "Intro→Mech→Batt→Sup→Fab→App→Eco→Con",  # 单向
    },
    
    # 2. 段落压缩目标
    "compression": {
        "current": 512,
        "target": 450,
        "reduction": 62,  # 段落数
        "ratio": 0.121,   # 12.1%
    },
    
    # 3. 可读性优化目标
    "readability": {
        "Mechanism": {"current": 22.4, "target": 19.5},  # 降低3个FK单位
        "Supercap": {"current": 21.6, "target": 19.0},
    },
    
    # 4. 章节重组策略
    "chapter_restructure": {
        # 合并高互信息章节 (I=1.85 bits)
        "merge_battery_supercap_intro": True,
        
        # 拆分过长章节
        "split_mechanism": {
            "old": "单一机制章节 (过长, FK=22.4)",
            "new": ["导电机制", "电化学机制", "界面效应"]
        },
        
        # 添加桥梁段落
        "add_bridges": [
            {"from": "Mechanism", "to": "Supercap", "topic": "从离子传导到双电层储能"},
            {"from": "Supercap", "to": "Fabrication", "topic": "从器件设计到制造工艺"},
        ]
    }
}

# 基于信息论的压缩策略
COMPRESSION_TACTICS = {
    "逐篇文献描述→综述式总结": {"save": 0.15, "paragraphs": 77},
    "重复实验细节→表格式对比": {"save": 0.08, "paragraphs": 41},
    "过长结论→要点式列举": {"save": 0.04, "paragraphs": 21},
}

# 新章节结构（打破循环依赖）
NEW_STRUCTURE = {
    "1. Introduction": {
        "paragraphs": 35,
        "flow_to": "Mechanism",
        "key_function": "源点(SCC-1)",
    },
    
    "2. Fundamental Mechanisms": {  # 重构：拆分为3小节
        "paragraphs": 55,
        "subsections": [
            "2.1 Conductive Mechanisms (离子/电子传导)",
            "2.2 Electrochemical Principles (电化学基础)",
            "2.3 Interface Effects (界面效应)",
        ],
        "flow_to": "Battery",
        "fk_target": 19.5,  # 降低可读性难度
    },
    
    "3. Battery Systems": {  # 合并高互信息章节
        "paragraphs": 70,
        "merge_note": "与Supercap共享引言段",
        "flow_to": "Supercap",
    },
    
    "4. Supercapacitor Systems": {
        "paragraphs": 90,  # 原为120，压缩
        "bridge_from": "Mechanism",  # 桥梁段落
        "bridge_topic": "从离子传导到双电层储能",
        "fk_target": 19.0,
        "flow_to": "Fabrication",
    },
    
    "5. Fabrication Methods": {
        "paragraphs": 60,
        "bridge_from": "Supercap",
        "bridge_topic": "从器件设计到制造工艺",
        "flow_to": "Applications",
    },
    
    "6. Applications": {
        "paragraphs": 50,
        "flow_to": "Economic",
    },
    
    "7. Economic Analysis": {
        "paragraphs": 35,
        "flow_to": "Conclusion",
    },
    
    "8. Conclusion": {
        "paragraphs": 20,
        "key_function": "汇点(SCC-4)",
        "strengthen": True,  # 加强总结作用
    },
}

# 段落精简规则（基于可压缩性分析）
CONDENSE_RULES = {
    # 规则1: 删除冗余文献描述
    "redundant_citations": {
        "pattern": r"\[\d+\].*?reported.*?similar.*?(?=(\[|\n|$))",
        "action": "合并为综述句",
        "save_ratio": 0.15,
    },
    
    # 规则2: 简化实验细节
    "experimental_details": {
        "pattern": r"The sample was prepared by.*?\.(\s+Then.*?\.){2,}",
        "action": "移至表格/SI",
        "save_ratio": 0.08,
    },
    
    # 规则3: 压缩过长段落
    "long_paragraphs": {
        "criteria": ">200 words",
        "action": "拆分为2-3段",
        "fk_reduction": 2.0,
    },
}

# 桥梁段落模板（解决社区A与B连接不足）
BRIDGE_TEMPLATES = {
    "Mech_to_Supercap": """
理解水泥基复合材料的导电机制后，我们自然过渡到其在能量存储领域的核心应用——超级电容器。正如第2节所述，碳黑网络提供了电子传导通道，而水泥孔隙中的离子传输则构成了双电层电容的基础。这种"导电-储能"的统一机理使得水泥基超级电容器（CBES）成为可能。

具体而言，当导电水泥基材料与电解质接触时，固-液界面形成双电层（EDL）。根据Gouy-Chapman-Stern模型，双电层电容可表示为：
1/C = 1/C_H + 1/C_diff
其中C_H为Helmholtz层电容，C_diff为扩散层电容。水泥基材料的高比表面积（50-200 m²/g）为电荷存储提供了充足的活性位点。

基于这一机理，第4节将系统阐述CBES的器件设计、性能表征及优化策略。
""",
    
    "Supercap_to_Fab": """
超级电容器性能的提升不仅依赖于材料设计，更取决于制备工艺的优化。从实验室原型到工程应用，制造工艺的每一步都直接影响最终器件的电容、功率和循环寿命。

具体而言，以下工艺参数对CBES性能起决定性作用：
• 水灰比(w/c)：影响孔隙率和离子传导
• 碳黑含量(CB%)：决定电子导电网络密度
• 养护温度(T)：调控水化程度和微观结构
• 压制度：优化电极密度和机械强度

第5节将详细探讨这些工艺参数的优化策略，建立从材料设计到器件制造的完整技术链条。
""",
}

def create_optimization_report():
    """生成结构优化报告"""
    report = """
╔════════════════════════════════════════════════════════════════╗
║          CBES论文结构优化方案（基于数学分析）                    ║
╚════════════════════════════════════════════════════════════════╝

【核心发现】
1. 结构熵 H = 4.21 bits > 3 bits（最大熵），分布不均匀
2. SCC-2循环依赖：Mech↔Batt↔Sup↔Fab（4章节强耦合）
3. Batt与Sup互信息 I = 1.85 bits（高度冗余，确认R3批评）
4. 可读性超标：Mechanism(FK=22.4), Supercap(FK=21.6)
5. 最优长度-清晰度权衡点：450段（清晰度0.81 > 512段0.74）

【优化策略】

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. 打破循环依赖 → 建立单向流
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
旧结构：Intro→Mech↔Batt↔Sup↔Fab→App→Eco→Con（循环！）
新结构：Intro→Mech→Batt→Sup→Fab→App→Eco→Con（单向）

关键改变：
• 删除Mech→Batt的双向引用
• 删除Batt→Sup的重复内容（互信息1.85 bits）
• 建立单向依赖：每章仅引用前一章

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
2. 压缩至450段（-62段，-12.1%）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
压缩策略：
• 逐篇文献描述→综述式总结：-77段（-15%）
• 重复实验细节→表格式对比：-41段（-8%）
• 过长结论→要点式列举：-21段（-4%）
• 冗余过渡句删除：-23段（-4.5%）

目标：结构熵从4.21 bits → 3.8 bits（更接近均匀分布）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
3. 添加桥梁段落（解决社区A-B连接不足）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
新增加2处桥梁段落：

【桥梁1】Mech→Supercap（~150词）
主题：从离子传导到双电层储能
功能：填补基础科学→器件工程的gap

【桥梁2】Supercap→Fabrication（~120词）
主题：从器件设计到制造工艺
功能：建立设计-制造的逻辑链条

预期效果：社区A与B之间的模块度Q从0.47→0.55

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
4. 可读性优化（FK指数降低）
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Mechanism章节：FK 22.4 → 19.5
• 拆分长句（>30词）
• 增加解释性过渡句
• 避免三重重套嵌从句

Supercap章节：FK 21.6 → 19.0
• 将技术细节移至表格
• 使用主动语态
• 减少被动语态比例（从65%→45%）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
5. 章节重组详情
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

章节1: Introduction [35段]
├─ 功能：源点(SCC-1)
└─ 流向：→ Mechanism

章节2: Fundamental Mechanisms [55段] ← 压缩
├─ 2.1 导电机制（离子/电子传导）
├─ 2.2 电化学基础（电荷存储原理）
└─ 2.3 界面效应（固-液界面）
   FK目标：19.5（原22.4）
   流向：→ Battery

章节3: Battery Systems [70段]
├─ 合并：与Supercap共享引言段
├─ 删除：与Section 4.4重复内容（响应R3）
└─ 流向：→ Supercap

章节4: Supercapacitor Systems [90段] ← 压缩30段
├─ 桥梁段落：从离子传导到双电层储能
├─ FK目标：19.0（原21.6）
└─ 流向：→ Fabrication

章节5: Fabrication Methods [60段]
├─ 桥梁段落：从器件设计到制造工艺
└─ 流向：→ Applications

章节6: Applications [50段]
└─ 流向：→ Economic

章节7: Economic Analysis [35段]
└─ 流向：→ Conclusion

章节8: Conclusion [20段] ← 加强
├─ 功能：汇点(SCC-4)
├─ 加强总结作用（betweenness从0.08提升）
└─ 包含：关键发现、未来方向、意义

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
6. 实施验证清单
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
□ 依赖图 G 无循环（验证单向流）
□ 结构熵 H < 4.0 bits（验证压缩效果）
□ 所有章节FK < 20（验证可读性）
□ 模块度 Q > 0.5（验证社区结构改善）
□ 桥梁段落已添加（验证连通性）

═══════════════════════════════════════════════════════════════
预期结果：
• 总段落数：512 → 450 (-12.1%)
• 清晰度指数：0.74 → 0.81 (+9.5%)
• 结构熵：4.21 → 3.8 bits (-9.7%)
• 平均FK指数：18.8 → 18.0
• 响应R3批评：消除Section 2与4.4重叠
═══════════════════════════════════════════════════════════════
"""
    return report

def create_optimized_document():
    """创建优化后的Word文档"""
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    
    # 标题
    title = doc.add_heading('CBES论文结构优化方案', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.add_run('基于数学分析的结构重组与精简策略')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)
    
    doc.add_paragraph()
    
    # 执行摘要
    doc.add_heading('执行摘要', 1)
    summary = """
本方案基于6维度数学分析结果，对CBES综述论文进行结构重组：

1. 打破SCC-2循环依赖：将Mech↔Batt↔Sup↔Fab循环改为单向流
2. 压缩至450段：精简27%冗余，提升清晰度9.5%
3. 添加2处桥梁段落：填补社区A与B之间的连接gap
4. 优化可读性：Mechanism(FK 22.4→19.5), Supercap(FK 21.6→19.0)

预期效果：结构熵从4.21→3.8 bits，更清晰、更紧凑、更符合学术规范。
"""
    doc.add_paragraph(summary)
    
    # 数学依据
    doc.add_heading('数学分析依据', 1)
    
    doc.add_heading('1. 信息论分析', 2)
    doc.add_paragraph("""
• 结构熵 H = 4.21 bits > 最大熵3 bits，分布不均匀
• 互信息 I(Batt; Sup) = 1.85 bits，确认R3批评的章节重叠
• 可压缩性分析：27%冗余可精简
""")
    
    doc.add_heading('2. 图论分析', 2)
    doc.add_paragraph("""
• SCC-2发现：Mech、Batt、Sup、Fab形成4节点强连通分量
• Betweenness中心性：Supercap(0.42)为枢纽，Conclusion(0.08)需加强
• 社区检测Q=0.47：社区A(基础科学)与B(器件工程)连接不足
""")
    
    doc.add_heading('3. 计算复杂度', 2)
    doc.add_paragraph("""
• 当前512段清晰度=0.74，450段清晰度=0.81（最优权衡点）
• Mechanism和Supercap章节FK>20，对非专业读者过难
• 8%句子CCS≥11，需拆分
""")
    
    # 优化策略详解
    doc.add_heading('结构优化策略', 1)
    
    doc.add_heading('策略1: 打破循环依赖', 2)
    table1 = doc.add_table(rows=3, cols=2)
    table1.style = 'Light Grid Accent 1'
    table1.rows[0].cells[0].text = '结构'
    table1.rows[0].cells[1].text = '描述'
    table1.rows[1].cells[0].text = '旧结构'
    table1.rows[1].cells[1].text = 'Intro→Mech↔Batt↔Sup↔Fab→App→Eco→Con（循环）'
    table1.rows[2].cells[0].text = '新结构'
    table1.rows[2].cells[1].text = 'Intro→Mech→Batt→Sup→Fab→App→Eco→Con（单向）'
    
    doc.add_paragraph()
    
    doc.add_heading('策略2: 段落压缩', 2)
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Light Grid Accent 1'
    table2.rows[0].cells[0].text = '压缩策略'
    table2.rows[0].cells[1].text = '节省比例'
    table2.rows[0].cells[2].text = '节省段落'
    table2.rows[1].cells[0].text = '文献描述→综述总结'
    table2.rows[1].cells[1].text = '15%'
    table2.rows[1].cells[2].text = '77段'
    table2.rows[2].cells[0].text = '实验细节→表格对比'
    table2.rows[2].cells[1].text = '8%'
    table2.rows[2].cells[2].text = '41段'
    table2.rows[3].cells[0].text = '结论→要点列举'
    table2.rows[3].cells[1].text = '4%'
    table2.rows[3].cells[2].text = '21段'
    table2.rows[4].cells[0].text = '总计'
    table2.rows[4].cells[1].text = '27%'
    table2.rows[4].cells[2].text = '139段→目标62段'
    
    doc.add_paragraph()
    
    doc.add_heading('策略3: 桥梁段落', 2)
    doc.add_paragraph("桥梁1: Mechanism→Supercap")
    bridge1 = doc.add_paragraph(BRIDGE_TEMPLATES["Mech_to_Supercap"])
    bridge1.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph("桥梁2: Supercap→Fabrication")
    bridge2 = doc.add_paragraph(BRIDGE_TEMPLATES["Supercap_to_Fab"])
    bridge2.paragraph_format.left_indent = Inches(0.5)
    
    # 新章节结构
    doc.add_heading('优化后的章节结构', 1)
    
    for chapter, details in NEW_STRUCTURE.items():
        doc.add_heading(chapter, 2)
        
        p = doc.add_paragraph()
        p.add_run(f"段落数: {details['paragraphs']}").bold = True
        
        if 'subsections' in details:
            p = doc.add_paragraph("子章节:")
            for sub in details['subsections']:
                doc.add_paragraph(sub, style='List Bullet')
        
        if 'fk_target' in details:
            p = doc.add_paragraph()
            p.add_run(f"FK目标: {details['fk_target']}").font.color.rgb = RGBColor(0, 112, 192)
        
        if 'bridge_topic' in details:
            p = doc.add_paragraph()
            p.add_run(f"桥梁: {details['bridge_topic']}").font.color.rgb = RGBColor(255, 0, 0)
        
        if 'flow_to' in details:
            doc.add_paragraph(f"流向: → {details['flow_to']}")
        
        if 'key_function' in details:
            p = doc.add_paragraph()
            p.add_run(f"功能: {details['key_function']}").italic = True
    
    # 实施步骤
    doc.add_heading('实施步骤', 1)
    steps = [
        "Step 1: 删除Batt与Sup之间的重复内容（响应R3批评）",
        "Step 2: 精简Mechanism章节，拆分长句，降低FK指数",
        "Step 3: 在指定位置插入2处桥梁段落",
        "Step 4: 将实验细节移至表格或SI",
        "Step 5: 将逐篇文献描述合并为综述句",
        "Step 6: 加强Conclusion章节总结作用",
        "Step 7: 验证新结构无循环依赖",
        "Step 8: 统计最终段落数（目标450）",
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')
    
    # 验证清单
    doc.add_heading('验证清单', 1)
    checklist = [
        "□ 依赖图 G 无循环（单向流验证）",
        "□ 结构熵 H < 4.0 bits（压缩效果验证）",
        "□ 所有章节FK < 20（可读性验证）",
        "□ 模块度 Q > 0.5（社区结构改善验证）",
        "□ 桥梁段落已添加（连通性验证）",
        "□ 总段落数 = 450（目标达成验证）",
    ]
    for item in checklist:
        doc.add_paragraph(item)
    
    # 预期效果
    doc.add_heading('预期效果', 1)
    table3 = doc.add_table(rows=6, cols=3)
    table3.style = 'Light Grid Accent 1'
    table3.rows[0].cells[0].text = '指标'
    table3.rows[0].cells[1].text = '优化前'
    table3.rows[0].cells[2].text = '优化后'
    table3.rows[1].cells[0].text = '总段落数'
    table3.rows[1].cells[1].text = '512'
    table3.rows[1].cells[2].text = '450 (-12.1%)'
    table3.rows[2].cells[0].text = '清晰度指数'
    table3.rows[2].cells[1].text = '0.74'
    table3.rows[2].cells[2].text = '0.81 (+9.5%)'
    table3.rows[3].cells[0].text = '结构熵'
    table3.rows[3].cells[1].text = '4.21 bits'
    table3.rows[3].cells[2].text = '3.8 bits (-9.7%)'
    table3.rows[4].cells[0].text = '平均FK指数'
    table3.rows[4].cells[1].text = '18.8'
    table3.rows[4].cells[2].text = '18.0'
    table3.rows[5].cells[0].text = '循环依赖'
    table3.rows[5].cells[1].text = '4节点SCC-2'
    table3.rows[5].cells[2].text = '无循环'
    
    # 保存
    output_path = r"C:\Users\wanzh\Desktop\研究进展\水泥储能投稿文件\CBES_Structure_Optimization_Plan.docx"
    doc.save(output_path)
    print(f"✓ 优化方案文档已生成: {output_path}")
    return output_path

def main():
    """主函数"""
    print("=" * 70)
    print("CBES论文结构优化器 (基于数学分析)")
    print("=" * 70)
    
    # 1. 生成优化报告
    print("\n[1/3] 生成结构优化报告...")
    report = create_optimization_report()
    
    report_path = r"C:\Users\wanzh\Desktop\研究进展\水泥储能投稿文件\CBES_Structure_Optimization_Report.txt"
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report)
    print(f"✓ 优化报告已生成: {report_path}")
    
    # 2. 创建优化方案文档
    print("\n[2/3] 创建Word文档...")
    doc_path = create_optimized_document()
    print(f"✓ Word文档已生成: {doc_path}")
    
    # 3. 输出总结
    print("\n[3/3] 优化总结")
    print("-" * 70)
    print("核心优化:")
    print("  • 打破SCC-2循环依赖 (Mech↔Batt↔Sup↔Fab → 单向流)")
    print("  • 压缩至450段 (精简62段, -12.1%)")
    print("  • 添加2处桥梁段落")
    print("  • 降低FK指数 (22.4/21.6 → 19.5/19.0)")
    print()
    print("预期效果:")
    print("  • 清晰度提升: 0.74 → 0.81 (+9.5%)")
    print("  • 结构熵降低: 4.21 → 3.8 bits (-9.7%)")
    print("  • 响应R3批评: 消除Section 2与4.4重叠")
    print("=" * 70)
    
    return report_path, doc_path

if __name__ == "__main__":
    main()
