"""
Scientific Skills Enhancement Report for CBES Paper
Using research-paper-writer + pymatgen standards
"""

from docx import Document
import os

def analyze_and_enhance():
    doc = Document('CBES_Final_Revised_Full.docx')
    
    report = []
    report.append("=" * 80)
    report.append("SCIENTIFIC SKILLS ENHANCEMENT REPORT")
    report.append("CBES (Cement-Based Energy Storage) Paper Analysis")
    report.append("=" * 80)
    report.append("")
    
    # 1. Basic Statistics
    report.append("## 1. DOCUMENT STATISTICS")
    report.append("-" * 40)
    total_paras = len(doc.paragraphs)
    total_tables = len(doc.tables)
    total_images = sum(1 for shape in doc.inline_shapes if hasattr(shape, 'type'))
    
    non_empty_paras = sum(1 for p in doc.paragraphs if p.text.strip())
    total_chars = sum(len(p.text) for p in doc.paragraphs if p.text.strip())
    
    report.append(f"Total Paragraphs: {total_paras}")
    report.append(f"Non-empty Paragraphs: {non_empty_paras}")
    report.append(f"Total Characters: {total_chars:,}")
    report.append(f"Total Tables: {total_tables}")
    report.append(f"Total Images: {total_images}")
    report.append("")
    
    # 2. Research Paper Writer Standards Check
    report.append("## 2. ACADEMIC WRITING STANDARDS (IEEE/ACM)")
    report.append("-" * 40)
    
    # Find Abstract
    abstract_text = ""
    abstract_found = False
    for para in doc.paragraphs:
        if 'Abstract' in para.text and len(para.text) > 100:
            abstract_text = para.text
            abstract_found = True
            break
    
    if abstract_found:
        word_count = len(abstract_text.split())
        status = "✓ PASS" if 150 <= word_count <= 250 else "⚠ NEEDS ADJUSTMENT"
        report.append(f"[{status}] Abstract Length: {word_count} words")
        report.append(f"         Target: 150-250 words (IEEE/ACM standard)")
        
        if word_count < 150:
            report.append("         → Suggestion: Expand abstract with more methodology details")
        elif word_count > 250:
            report.append("         → Suggestion: Condense to focus on key contributions")
    else:
        report.append("[⚠] Abstract not found or too short")
    
    report.append("")
    
    # Check for key sections
    required_sections = [
        'Introduction',
        'Conduction Mechanism', 
        'Structural Battery',
        'Structural Supercapacitor',
        'Safety',
        'Application',
        'Techno-Economic',
        'Conclusion',
        'Reference'
    ]
    
    report.append("Section Coverage:")
    section_found = {}
    for req_section in required_sections:
        found = any(req_section.lower() in p.text.lower() for p in doc.paragraphs)
        section_found[req_section] = found
        status = "✓" if found else "✗"
        report.append(f"  [{status}] {req_section}")
    
    missing_sections = [s for s, f in section_found.items() if not f]
    if missing_sections:
        report.append(f"\nMissing Sections: {', '.join(missing_sections)}")
    
    report.append("")
    
    # 3. Pymatgen Materials Science Analysis
    report.append("## 3. MATERIALS SCIENCE ANALYSIS (Pymatgen Standards)")
    report.append("-" * 40)
    
    # Analyze material-related content
    materials_keywords = {
        'Cement-based': ['cement', 'concrete', 'C-S-H', 'pore solution', 'hydration'],
        'Electrode': ['carbon fiber', 'CNT', 'graphene', 'aerogel', 'conductive polymer'],
        'Electrolyte': ['KOH', 'NaCl', 'H2SO4', 'gel polymer', 'ionic liquid'],
        'Properties': ['conductivity', 'capacitance', 'impedance', 'mechanical strength']
    }
    
    keyword_counts = {}
    for category, keywords in materials_keywords.items():
        count = 0
        for kw in keywords:
            count += sum(1 for p in doc.paragraphs if kw.lower() in p.text.lower())
        keyword_counts[category] = count
    
    report.append("Material Keywords Frequency:")
    for category, count in sorted(keyword_counts.items(), key=lambda x: x[1], reverse=True):
        bar = '█' * min(count // 5, 20)
        report.append(f"  {category:20s}: {count:4d} {bar}")
    
    report.append("")
    
    # 4. Figure Quality Assessment
    report.append("## 4. FIGURE QUALITY ASSESSMENT")
    report.append("-" * 40)
    
    expected_figures = [
        ('Figure 1', 'Development Roadmap'),
        ('Figure 2', 'Technology Comparison'),
        ('Figure 3', 'Safety/Lifecycle (Combined)'),
        ('Figure 4', 'Application Framework'),
        ('Figure 5', 'Electrode Performance'),
        ('Figure 6', 'Electrolyte System'),
        ('Figure 7', 'Device Architecture'),
        ('Figure 8', 'Bibliometric Analysis'),
    ]
    
    figure_files = [f for f in os.listdir('.') if f.startswith('Figure') and f.endswith('.png')]
    
    report.append(f"Expected Figures: {len(expected_figures)}")
    report.append(f"Actual Figure Files: {len(figure_files)}")
    report.append("")
    
    for fig_name, desc in expected_figures:
        exists = any(fig_name.replace(' ', '_') in f or fig_name in f for f in figure_files)
        status = "✓ Found" if exists else "✗ Missing"
        report.append(f"  [{status}] {fig_name}: {desc}")
    
    report.append("")
    
    # 5. Reviewer Response Coverage
    report.append("## 5. REVIEWER RESPONSE COVERAGE")
    report.append("-" * 40)
    
    reviewer_concerns = {
        'Reviewer 1 (Structure)': {
            'Fragmented narrative': ['roadmap', 'clear logic chain'],
            'Service life mismatch': ['modular design', 'predictive monitoring', 'redundant config'],
            'Safety concerns': ['leakage', 'encapsulation', 'thermal management'],
            'Self-discharge': ['self-discharge rate', 'comparison with Li-ion'],
            'Volume/weight advantage': ['specific energy', 'power density comparison'],
        },
        'Reviewer 2 (Minor)': {
            'Long sentences': ['sentence length < 30 words'],
            'Small fonts in figures': ['font size ≥ 8pt at 300 DPI'],
            'Definitions unclear': ['structural battery vs supercapacitor definition'],
            'Economic assumptions': ['sensitivity analysis', 'uncertainty discussion'],
            'Reference format': ['consistent formatting', 'DOI included'],
        },
        'Reviewer 3 (Major)': {
            'Redundant content': ['merged conduction mechanisms', 'removed duplicates'],
            'Lack of critical analysis': ['critical comparison tables', 'standardized metrics'],
            'Loose Section 6': ['unified framework', 'TRL organization'],
            'Figure quality': ['removed schematic diagrams', 'data-focused figures'],
            'Language quality': ['native English editing', 'professional polish'],
            'Irrelevant citations': ['removed citation [262]', 'relevant references only'],
        }
    }
    
    for reviewer, concerns in reviewer_concerns.items():
        report.append(f"\n{reviewer}:")
        for concern, solutions in concerns.items():
            addressed = all(any(sol in p.text.lower() for p in doc.paragraphs) 
                         for sol in solutions)
            status = "✓ Addressed" if addressed else "⚠ Partial"
            report.append(f"  [{status}] {concern}")
    
    report.append("")
    
    # 6. Recommendations
    report.append("## 6. ENHANCEMENT RECOMMENDATIONS")
    report.append("-" * 40)
    
    recommendations = [
        ("HIGH", "Add standardized testing protocol section", 
         "Define specific test conditions (electrode thickness, voltage window, loading rate)"),
        ("HIGH", "Include crystal structure analysis of C-S-H gel",
         "Use pymatgen to analyze cement hydration products structure"),
        ("MEDIUM", "Create performance normalization table",
         "Normalize all reported values to standard conditions for fair comparison"),
        ("MEDIUM", "Add uncertainty quantification to economic model",
         "Include Monte Carlo simulation or confidence intervals"),
        ("LOW", "Consider adding supplementary information file",
         "Include raw data, additional figures, or extended methods"),
    ]
    
    for priority, title, desc in recommendations:
        report.append(f"\n[{priority}] {title}")
        report.append(f"     {desc}")
    
    report.append("")
    report.append("=" * 80)
    report.append("END OF ANALYSIS REPORT")
    report.append("=" * 80)
    
    # Save report
    output_file = 'Scientific_Skills_Analysis_Report.txt'
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))
    
    print(f'\n{"="*60}')
    print(f'✓ Scientific Skills Analysis Complete!')
    print(f'📄 Report saved: {output_file}')
    print(f'{"="*60}')
    
    return '\n'.join(report)

if __name__ == '__main__':
    print(analyze_and_enhance())
